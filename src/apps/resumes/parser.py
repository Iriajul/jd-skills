from io import BytesIO

from docx import Document
from pypdf import PdfReader


def extract_text_from_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(pages).strip()


def extract_text_from_docx(content: bytes) -> str:
    doc = Document(BytesIO(content))
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    # Include table cell text (some resumes lay out contact info in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    lines.append(cell.text)
    return "\n".join(lines).strip()


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> list[str]:
    """Split text into overlapping chunks suitable for embedding."""
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start += chunk_size - overlap
    return chunks
