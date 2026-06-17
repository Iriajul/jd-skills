"""Read paragraphs from a .docx and rewrite them in place, preserving formatting.

The rewrite keeps each paragraph's style and its first run's formatting (font,
size, bold, bullet style) and only swaps the text — so the tailored document is
visually identical to the original except for the wording.
"""
from io import BytesIO

from docx import Document


def list_paragraphs(content: bytes) -> list[tuple[int, str]]:
    """Return (position_index, text) for every non-empty paragraph.

    The index is the paragraph's absolute position in the document so it stays
    stable between reading (tailor) and writing (render).
    """
    doc = Document(BytesIO(content))
    out: list[tuple[int, str]] = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            out.append((i, text))
    return out


def _rewrite_paragraph(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        return
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def render(content: bytes, edits: dict[int, str]) -> bytes:
    """Apply {position_index: new_text} edits to the original docx; return new bytes."""
    doc = Document(BytesIO(content))
    paras = doc.paragraphs
    for index, new_text in edits.items():
        if 0 <= index < len(paras):
            _rewrite_paragraph(paras[index], new_text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
